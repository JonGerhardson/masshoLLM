import requests
import logging
import dateparser
import pypdf
import io
import pandas as pd
import xlrd # For .xls files
import json
from bs4 import BeautifulSoup
from trafilatura import extract
from typing import Optional, Dict, Any, List
from datetime import datetime, timedelta

# --- New Imports for Robustness ---
import cloudscraper
from fake_useragent import UserAgent

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Custom Exception for 403 Errors ---
class Scraper403Error(Exception):
    """Custom exception for 403 Forbidden errors to handle them specifically."""
    pass

# --- Global UserAgent Instance ---
ua = UserAgent()

def create_session_with_retries() -> requests.Session:
    """Creates a cloudscraper session object with more robust headers to bypass anti-bot measures."""
    session = cloudscraper.create_scraper()
    # ---  default headers ---
    session.headers.update({
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'Accept-Language': 'en-US,en;q=0.9',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
        'Cache-Control': 'max-age=0',
    })
    return session

def fetch_page(session: requests.Session, url: str) -> Optional[str]:
    """Fetches the HTML content of a given URL using a robust session object."""
    try:
        # --- UPGRADE: Add a Referer header to make the request look more natural ---
        headers = {
            'User-Agent': ua.random,
            'Referer': 'https://www.mass.gov/'
        }
        response = session.get(url, headers=headers, timeout=20)
        
        if response.status_code == 403:
            raise Scraper403Error(f"403 Forbidden error for URL: {url}")
            
        response.raise_for_status()
        return response.text
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching URL {url}: {e}")
        return None

def extract_download_page_date(html: str) -> Optional[datetime]:
    """Parses the HTML of a /doc/ landing page to find the 'Last updated' date."""
    try:
        soup = BeautifulSoup(html, 'html.parser')
        last_updated_th = soup.find('th', scope="row", string=lambda t: t and "last updated" in t.lower())
        if last_updated_th:
            next_td = last_updated_th.find_next_sibling('td')
            if next_td:
                date_span = next_td.find('span', class_='ma__listing-table__data-item')
                if date_span:
                    return dateparser.parse(date_span.get_text(strip=True))
        
        last_updated_dt = soup.find('dt', string=lambda t: t and "last updated on" in t.lower())
        if last_updated_dt:
            last_updated_dd = last_updated_dt.find_next_sibling('dd')
            if last_updated_dd:
                return dateparser.parse(last_updated_dd.get_text(strip=True))
        return None
    except Exception as e:
        logging.error(f"Error parsing date from download page: {e}")
        return None

def find_best_date_on_page(html: str) -> Dict[str, Optional[datetime]]:
    """
    Tries to find any plausible date from a standard HTML page using multiple methods.
    
    UPDATE: Now returns a dictionary distinguishing between a 'posting_date' 
    (for articles/updates) and a 'meeting_date' (for events).
    """
    results = {
        'posting_date': None,
        'meeting_date': None
    }
    try:
        soup = BeautifulSoup(html, 'html.parser')

        # --- NEW: Prioritize finding a specific meeting/event date ---
        # This is a common selector on Mass.gov event pages
        event_time_tag = soup.find('time', class_='ma-page-header__event-date', attrs={'datetime': True})
        if event_time_tag:
            results['meeting_date'] = dateparser.parse(event_time_tag['datetime'])
            # If we find a meeting date, we can often stop, as this is the most specific date.
            # However, we'll continue just in case a posting date is also relevant.

        # --- Standard Posting Date Logic ---
        meta_tag = soup.find('meta', property='article:published_time')
        if meta_tag and meta_tag.get('content'):
            results['posting_date'] = dateparser.parse(meta_tag['content'])
        
        press_date_div = soup.find('div', class_='ma__press-status__date')
        if press_date_div:
            results['posting_date'] = dateparser.parse(press_date_div.get_text(strip=True))
            
        date_span = soup.find('span', class_='ma-page-header__published-date')
        if date_span:
            results['posting_date'] = dateparser.parse(date_span.get_text(strip=True).replace("Published on ", ""))
            
        time_tag = soup.find('time', attrs={'datetime': True})
        if time_tag and not results['posting_date'] and not results['meeting_date']:
            # Use this as a fallback posting date if nothing else was found
            results['posting_date'] = dateparser.parse(time_tag['datetime'])
            
        return results
    except Exception as e:
        logging.error(f"Error finding a generic date on page: {e}")
        return results

def extract_article_content(html: str) -> Optional[str]:
    """Uses 'trafilatura' to extract main article content with quality verification."""
    # First, try the primary extraction method
    extracted_content = extract(html, include_comments=False, include_tables=False)
    
    if extracted_content:
        # Verify extraction quality
        if is_content_quality_acceptable(extracted_content, html):
            return extracted_content
    
    # If primary extraction fails quality check, try alternate methods
    alt_extracted = extract_with_alternate_methods(html)
    if alt_extracted and is_content_quality_acceptable(alt_extracted, html):
        return alt_extracted
    
    # If all methods fail, return the best available
    return extracted_content

def is_content_quality_acceptable(extracted_content: str, original_html: str) -> bool:
    """Verify that extracted content meets quality standards."""
    if not extracted_content or len(extracted_content.strip()) < 50:
        return False
    
    # Check for common extraction problems
    if extracted_content.count('\n') > len(extracted_content.split()) * 0.5:  # Too many line breaks
        return False
    
    # Check for repeated characters (common extraction artifacts)
    if extracted_content.count('  ') > len(extracted_content) * 0.05:  # Too many double spaces
        return False
    
    # If content is more than 90% of original HTML, it might be poorly extracted
    if len(extracted_content) > len(original_html) * 0.8:
        return False
    
    return True

def extract_with_alternate_methods(html: str) -> Optional[str]:
    """Try alternate content extraction methods."""
    soup = BeautifulSoup(html, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
        script.decompose()
    
    # Try different content selectors commonly found on mass.gov
    selectors_to_try = [
        'main',  # Main content area
        '.ma__content-main',  # Mass.gov main content class
        '.ma__body-content',  # Mass.gov body content class
        '.content',  # Generic content class
        'article',  # Article tag
        '.main-content',  # Main content class
        '.post-content',  # Post content class
        'body'  # Body tag
    ]
    
    for selector in selectors_to_try:
        elements = soup.select(selector)
        if elements:
            content = ' '.join([elem.get_text(separator=' ', strip=True) for elem in elements])
            if content and len(content) > 100:  # Only return if substantial content found
                return content
    
    # Fallback: get all text from body
    body = soup.find('body')
    if body:
        return body.get_text(separator=' ', strip=True)
    
    # Last resort: return original HTML text
    return soup.get_text(separator=' ', strip=True)

def download_and_extract_document_text(session: requests.Session, url: str) -> Dict[str, Optional[str]]:
    """
    Downloads a file, identifies its type based on content, and extracts all text from it.
    This function includes quality checks to ensure accurate extraction.
    """
    try:
        # --- UPGRADE: Add a Referer header to the download request as well ---
        headers = {
            'User-Agent': ua.random,
            'Referer': 'https://www.mass.gov/'
            }
        response = session.get(url, headers=headers, timeout=30)
        
        if response.status_code == 403:
            raise Scraper403Error(f"403 Forbidden error for URL: {url}")
            
        response.raise_for_status()
        content_bytes = response.content
        file_stream = io.BytesIO(content_bytes)

        # Identify file type by magic number for accuracy
        magic_number = content_bytes[:4]
        text = None
        file_type = "Unknown"

        if magic_number == b'%PDF':
            file_type = "PDF"
            text = extract_pdf_content(file_stream, url)
        elif magic_number == b'PK\x03\x04': # ZIP archive (DOCX, XLSX)
            try:
                file_type = "XLSX"
                df = pd.read_excel(file_stream, sheet_name=None, engine='openpyxl')
                text = "\n\n".join(f"Sheet: {name}\n{sheet.to_string()}" for name, sheet in df.items())
            except Exception:
                try:
                    import docx
                    file_type = "DOCX"
                    text = extract_docx_content(file_stream, url)
                except Exception as e:
                    logging.warning(f"File at {url} is a ZIP archive but not a readable XLSX or DOCX: {e}")
        elif magic_number == b'\xd0\xcf\x11\xe0': # Legacy XLS file
             file_type = "XLS"
             workbook = xlrd.open_workbook(file_contents=content_bytes)
             text = ""
             for sheet in workbook.sheets():
                 text += f"Sheet: {sheet.name}\n"
                 for row_idx in range(sheet.nrows):
                     text += "\t".join(str(cell.value) for cell in sheet.row(row_idx)) + "\n"
        else: # Fallback to CSV/plain text
            try:
                file_type = "CSV"
                text = pd.read_csv(file_stream).to_string()
            except Exception as e:
                logging.error(f"Could not read file as CSV: {e}")

        # Verify content quality before returning
        if text:
            # Clean up the extracted text
            text = clean_extracted_text(text)
            # Verify the content quality
            if is_document_content_quality_acceptable(text):
                return {"content_for_llm": text, "filetype": file_type}
            else:
                logging.warning(f"Content quality check failed for {url}. Attempting extraction with alternative methods...")
                # Try alternative methods if quality check fails
                alt_text = extract_with_alternative_methods(content_bytes, magic_number, file_stream, url)
                if alt_text and is_document_content_quality_acceptable(alt_text):
                    return {"content_for_llm": alt_text, "filetype": file_type}
                else:
                    logging.warning(f"All extraction methods failed quality check for {url}")

        return {"content_for_llm": text.strip() if text else None, "filetype": file_type}

    except pypdf.errors.PdfReadError:
        logging.warning(f"File at {url} appears to be a PDF but is unreadable.")
        return {"content_for_llm": None, "filetype": "PDF (unreadable)"}
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to download file from {url}: {e}")
        return {"content_for_llm": None, "filetype": "Download Failed"}
    except Exception as e:
        logging.error(f"An unexpected error occurred during file processing for {url}: {e}", exc_info=True)
        return {"content_for_llm": None, "filetype": "Processing Error"}

def extract_pdf_content(file_stream, url: str) -> str:
    """Extract text from PDF with quality improvements."""
    try:
        pdf_reader = pypdf.PdfReader(file_stream)
        text_pages = []
        
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_pages.append(page_text)
        
        text = "\n\n".join(text_pages)
        return text
    except Exception as e:
        logging.warning(f"Error extracting PDF content from {url}: {e}")
        return ""

def extract_docx_content(file_stream, url: str) -> str:
    """Extract text from DOCX files with quality improvements."""
    try:
        import docx
        doc = docx.Document(file_stream)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and cell.text not in paragraphs:
                        paragraphs.append(cell.text)
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        logging.warning(f"Error extracting DOCX content from {url}: {e}")
        return ""

def is_document_content_quality_acceptable(text: str) -> bool:
    """Check if extracted document content meets quality standards."""
    if not text or len(text.strip()) < 50:
        return False
    
    # Basic quality checks
    lines = text.split('\n')
    if len(lines) > 0:
        avg_line_length = sum(len(line) for line in lines) / len(lines)
        if avg_line_length < 10:  # Lines too short, possible extraction issue
            return False
    
    # Check for repeated content which might indicate extraction artifacts
    text_lines = text.split('\n')
    unique_lines = set(line.strip() for line in text_lines if line.strip())
    if len(unique_lines) < len(text_lines) * 0.3:  # More than 70% duplicate lines
        return False
    
    return True

def extract_with_alternative_methods(content_bytes, magic_number, file_stream, url: str) -> str:
    """Try alternative extraction methods for different file types."""
    # Reset file stream position
    file_stream.seek(0)
    
    if magic_number == b'%PDF':
        # Try alternative PDF extraction
        try:
            # Try different approach - extract text differently
            pdf_reader = pypdf.PdfReader(file_stream)
            full_text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        # Clean up the text
                        text = ' '.join(text.split())  # Normalize whitespace
                        full_text += text + " "
                except Exception as e:
                    logging.warning(f"Error extracting text from page {page_num} of {url}: {e}")
                    continue
            return full_text.strip()
        except Exception:
            pass
    
    elif magic_number == b'PK\x03\x04':
        # For DOCX, try different method
        try:
            import docx
            file_stream.seek(0)  # Reset position
            doc = docx.Document(file_stream)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text.strip())
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text.strip())
            
            return "\n\n".join([part for part in text_parts if part])
        except Exception:
            pass
    
    # If all alternatives fail, return empty string
    return ""

def clean_extracted_text(text: str) -> str:
    """Clean up extracted text to remove common artifacts."""
    if not text:
        return text
    
    # Remove extra whitespace
    import re
    text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespace with single space
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
    text = text.strip()
    
    return text

def fetch_press_releases(session: requests.Session) -> List[Dict[str, Any]]:
    """Fetches the latest press releases from the Mass.gov news API."""
    url = "https://www.mass.gov/api/v1/news"
    logging.info(f"Fetching press releases from {url}")
    try:
        # --- UPGRADE: Add browser-like headers to the API request to avoid 403 errors. ---
        headers = {
            'User-Agent': ua.random,
            'Referer': 'https://www.mass.gov/'
        }
        response = session.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        data = response.json()
        if isinstance(data, list):
            return data
        else:
            logging.error("Press release API response is not in the expected list format.")
            return []
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching press releases from API: {e}")
        return []
    except json.JSONDecodeError as e:
        logging.error(f"Error decoding JSON from press release API: {e}")
        return []

